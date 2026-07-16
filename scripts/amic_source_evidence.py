from __future__ import annotations

from dataclasses import dataclass
from email import policy
from email.parser import BytesParser
from pathlib import Path
from typing import Literal
from xml.etree import ElementTree
from zipfile import BadZipFile, ZipFile

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.exceptions import InvalidFileException
from pydantic import BaseModel, ConfigDict
from pypdf import PdfReader
from pypdf.errors import PdfReadError, PdfStreamError


ProfileKind = Literal[
    "civil_litigation",
    "criminal_litigation",
    "administrative_litigation",
    "corporate_advisory",
    "deal",
    "unknown",
]
ExtractionStatus = Literal["extracted", "empty", "skipped", "error"]
SUPPORTED_EXTENSIONS = frozenset({".csv", ".docx", ".eml", ".hwpx", ".pdf", ".rtf", ".txt", ".xlsx", ".xlsm"})


class SourceRecord(BaseModel):
    """Validated private source-index row used by the evidence extractor."""

    model_config = ConfigDict(extra="forbid", frozen=True)

    source_record_id: str
    relative_path: str
    source_scope: Literal["current", "legacy", "operations"]
    source_lane: str | None
    case_folder_hint: str | None
    extension: str
    file_kind: str
    extractability: str
    document_kind_hint: str
    profile_kind_hint: ProfileKind
    lane_mismatch_review_required: bool
    byte_size: int
    modified_at: str
    content_hash: str | None
    content_hash_status: str


class ExtractionOutcome(BaseModel):
    """Text extraction result that never persists the source document body."""

    model_config = ConfigDict(frozen=True)

    status: ExtractionStatus
    extractor: str
    text: str
    page_count: int | None
    truncated: bool
    error_kind: str | None


@dataclass(frozen=True, slots=True)
class ExtractionContext:
    """Immutable bounds for one source extraction run."""

    source_root: Path
    character_limit: int


def source_path(record: SourceRecord, context: ExtractionContext) -> Path:
    """Resolve a manifest-relative path without permitting directory traversal."""

    root = context.source_root.resolve()
    candidate = (root / record.relative_path).resolve()
    if root != candidate and root not in candidate.parents:
        raise ValueError("source record resolves outside the declared root")
    return candidate


def clip_text(text: str, character_limit: int) -> tuple[str, bool]:
    """Cap extracted content while preserving whether the source exceeded the cap."""

    return (text[:character_limit], len(text) > character_limit)


def document_text(path: Path, character_limit: int) -> ExtractionOutcome:
    """Extract visible DOCX paragraphs and table text."""

    document = Document(path)
    paragraph_text = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_text = [cell.text for table in document.tables for row in table.rows for cell in row.cells if cell.text.strip()]
    text, truncated = clip_text("\n".join([*paragraph_text, *table_text]), character_limit)
    return ExtractionOutcome(status="extracted" if text else "empty", extractor="python-docx", text=text, page_count=None, truncated=truncated, error_kind=None)


def hwpx_text(path: Path, character_limit: int) -> ExtractionOutcome:
    """Extract visible text nodes from HWPX XML parts."""

    with ZipFile(path) as archive:
        parts = sorted(name for name in archive.namelist() if name.startswith("Contents/") and name.endswith(".xml"))
        text = "\n".join("".join(ElementTree.fromstring(archive.read(part)).itertext()) for part in parts)
    clipped, truncated = clip_text(text, character_limit)
    return ExtractionOutcome(status="extracted" if clipped else "empty", extractor="hwpx-xml", text=clipped, page_count=None, truncated=truncated, error_kind=None)


def pdf_text(path: Path, character_limit: int) -> ExtractionOutcome:
    """Extract PDF page text up to the configured character bound."""

    reader = PdfReader(path)
    chunks: list[str] = []
    character_count = 0
    for page in reader.pages:
        page_text = page.extract_text() or ""
        chunks.append(page_text)
        character_count += len(page_text)
        if character_count >= character_limit:
            break
    text, truncated = clip_text("\n".join(chunks), character_limit)
    return ExtractionOutcome(status="extracted" if text else "empty", extractor="pypdf", text=text, page_count=len(reader.pages), truncated=truncated, error_kind=None)


def workbook_text(path: Path, character_limit: int) -> ExtractionOutcome:
    """Read cell values only; formulas, formatting, and workbook metadata are excluded."""

    workbook = load_workbook(path, read_only=True, data_only=True)
    try:
        chunks: list[str] = []
        character_count = 0
        for worksheet in workbook.worksheets:
            for row in worksheet.iter_rows(values_only=True):
                line = " | ".join(str(value) for value in row if value is not None).strip()
                if not line:
                    continue
                chunks.append(line)
                character_count += len(line) + 1
                if character_count >= character_limit:
                    text, _ = clip_text("\n".join(chunks), character_limit)
                    return ExtractionOutcome(status="extracted" if text else "empty", extractor="openpyxl", text=text, page_count=None, truncated=True, error_kind=None)
        text, truncated = clip_text("\n".join(chunks), character_limit)
        return ExtractionOutcome(status="extracted" if text else "empty", extractor="openpyxl", text=text, page_count=None, truncated=truncated, error_kind=None)
    finally:
        workbook.close()


def email_text(path: Path, character_limit: int) -> ExtractionOutcome:
    """Extract plain-text email parts and omit attachments and headers."""

    message = BytesParser(policy=policy.default).parsebytes(path.read_bytes())
    parts = [part.get_content() for part in message.walk() if part.get_content_type() == "text/plain" and not part.get_filename()]
    text, truncated = clip_text("\n".join(str(part) for part in parts), character_limit)
    return ExtractionOutcome(status="extracted" if text else "empty", extractor="email", text=text, page_count=None, truncated=truncated, error_kind=None)


def plain_text(path: Path, character_limit: int) -> ExtractionOutcome:
    """Read plain-text and CSV-compatible sources with replacement for invalid bytes."""

    text, truncated = clip_text(path.read_text(encoding="utf-8", errors="replace"), character_limit)
    return ExtractionOutcome(status="extracted" if text else "empty", extractor="plain-text", text=text, page_count=None, truncated=truncated, error_kind=None)


def extract_text(record: SourceRecord, context: ExtractionContext) -> ExtractionOutcome:
    """Route one supported source to its minimum viable text extractor."""

    if record.extension not in SUPPORTED_EXTENSIONS:
        return ExtractionOutcome(status="skipped", extractor="unsupported", text="", page_count=None, truncated=False, error_kind=None)
    try:
        path = source_path(record, context)
        match record.extension:
            case ".docx":
                return document_text(path, context.character_limit)
            case ".hwpx":
                return hwpx_text(path, context.character_limit)
            case ".pdf":
                return pdf_text(path, context.character_limit)
            case ".xlsx" | ".xlsm":
                return workbook_text(path, context.character_limit)
            case ".eml":
                return email_text(path, context.character_limit)
            case ".csv" | ".rtf" | ".txt":
                return plain_text(path, context.character_limit)
            case unreachable:
                raise ValueError(f"unsupported extension: {unreachable}")
    except (BadZipFile, ElementTree.ParseError, InvalidFileException, KeyError, OSError, PdfReadError, PdfStreamError, TypeError, UnicodeDecodeError, ValueError) as error:
        return ExtractionOutcome(status="error", extractor="error", text="", page_count=None, truncated=False, error_kind=type(error).__name__)
